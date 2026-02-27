import requests
import json
import io
import fitz

from pathlib import Path
from langchain_text_splitters import RecursiveCharacterTextSplitter
from typing import List
from langchain_core.documents import Document
from datetime import datetime, timedelta
from docx import Document as DocxDocument

from infra.settings import settings

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1500,
    chunk_overlap=150,
    separators=["\n\n", "\n", ". ", " ", ""],
)


class O365TokenManager:
    """Verwaltet OAuth-Tokens mit Client Credentials Flow für Single Tenant"""

    def __init__(self, token_file: Path = None):
        if token_file is None:
            token_dir = Path.home() / ".credentials"
            token_dir.mkdir(parents=True, exist_ok=True)
            token_file = token_dir / "o365_token.json"

        self.token_file = token_file

    def get_access_token(self) -> str:
        """Holt Access Token mit Client Credentials Flow"""

        if self.token_file.exists():
            try:
                with open(self.token_file, "r") as f:
                    token_data = json.load(f)

                if "expires_at" in token_data:
                    expires_at = datetime.fromisoformat(token_data["expires_at"])
                    if datetime.now() < expires_at:
                        print("✅ Token ist noch gültig")
                        return token_data["access_token"]
            except (json.JSONDecodeError, FileNotFoundError):
                pass

        print("🔄 Hole neuen Access Token mit Client Credentials Flow....")

        token_url = f"https://login.microsoftonline.com/{settings.SP_TENANT_ID}/oauth2/v2.0/token"

        payload = {
            "grant_type": "client_credentials",
            "client_id": settings.SP_CLIENT_ID,
            "client_secret": settings.SP_CLIENT_SECRET,
            "scope": "https://graph.microsoft.com/.default",
        }

        try:
            response = requests.post(token_url, data=payload, timeout=10)
            response.raise_for_status()
        except requests.exceptions.RequestException as e:
            print(f"❌ Fehler beim Token-Request: {e}")
            raise

        token_data = response.json()

        if "error" in token_data:
            print(f"❌ Token Error: {token_data.get('error_description', token_data['error'])}")
            raise Exception(f"Token Request failed: {token_data['error']}")

        expires_in = token_data.get("expires_in", 3600)
        expires_at = datetime.now() + timedelta(seconds=expires_in)

        token_cache = {
            "access_token": token_data["access_token"],
            "expires_at": expires_at.isoformat(),
            "expires_in": expires_in,
        }

        with open(self.token_file, "w") as f:
            json.dump(token_cache, f)

        print(f"✅ Token erfolgreich geholt und gespeichert: {self.token_file}")
        return token_data["access_token"]


class SharePointDocumentLoader:
    """Lädt Dokumente direkt von SharePoint über Microsoft Graph API"""

    def __init__(self, tenant_id: str, client_id: str, client_secret: str, site_id: str, drive_id: str):
        self.tenant_id = tenant_id
        self.client_id = client_id
        self.client_secret = client_secret
        self.site_id = site_id
        self.drive_id = drive_id

        self.token_manager = O365TokenManager()
        self.graph_api_base = "https://graph.microsoft.com/v1.0"
        self.headers = {}
        self._refresh_headers()

    def _refresh_headers(self):
        """Headers mit aktuellem Token aktualisieren"""
        access_token = self.token_manager.get_access_token()
        self.headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json",
        }

    def _fetch_files_recursive(self, folder_item_id: str = "root", path: str = "") -> List[dict]:
        """Rekursiv alle Dateien aus einem Ordner holen"""

        url = f"{self.graph_api_base}/drives/{self.drive_id}/items/{folder_item_id}/children"
        files = []

        try:
            response = requests.get(url, headers=self.headers, timeout=10)
            response.raise_for_status()

            items = response.json().get("value", [])

            for item in items:
                current_path = f"{path}/{item['name']}" if path else item['name']

                if item.get("folder"):
                    files.extend(self._fetch_files_recursive(item["id"], current_path))
                elif item.get("file"):
                    files.append({
                        "id": item["id"],
                        "name": item["name"],
                        "path": current_path,
                        "size": item.get("size", 0),
                        "download_url": item.get("@microsoft.graph.downloadUrl"),
                    })

        except requests.exceptions.RequestException as e:
            print(f"❌ Fehler beim Abrufen von Dateien: {e}")

        return files

    def _download_file_content(self, download_url: str, file_name: str) -> str:
        """Lädt Dateiinhalt herunter und extrahiert Text aus PDF/DOCX/TXT/MD."""
        try:
            if not any(file_name.lower().endswith(ext) for ext in [".txt", ".pdf", ".md", ".docx"]):
                print(f"⏭️ Überspringe {file_name} (nicht unterstützter Typ)")
                return ""

            response = requests.get(download_url, timeout=60)
            response.raise_for_status()

            # PDF -> Text
            if file_name.lower().endswith(".pdf"):
                pdf_bytes = response.content
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text

            # DOCX -> Text
            if file_name.lower().endswith(".docx"):
                docx_file = io.BytesIO(response.content)
                doc = DocxDocument(docx_file)
                return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

            # TXT/MD -> Text
            return response.text

        except Exception as e:
            print(f"❌ Fehler beim Download/Parsing ({file_name}): {e}")
            return ""

    def load_documents(self, folder_path: str | None = None) -> List[Document]:
        """Lädt alle Dokumente aus SharePoint (optional nur aus bestimmten Ordnern)"""

        print(f"📂 Lade Dokumente von SharePoint Drive: {self.drive_id}")
        files = self._fetch_files_recursive()
        print(f"✅ {len(files)} Dateien gefunden")

        # Nur Dateien aus diesen Ordnern laden
        required_subpaths = [
            "Allgemein/01_Unternehmensdokumente",
        ]

        documents = []

        for file in files:
            path = file["path"]

            # Nur Dateien aus den erforderlichen Ordnern laden
            if not any(sub in path for sub in required_subpaths):
                continue

            print(f"  📄 Verarbeite: {path} ({file['size']} bytes)")

            content = self._download_file_content(file["download_url"], file["name"])
            if not content:
                continue

            doc = Document(
                page_content=content,
                metadata={
                    "source": path,
                    "file_id": file["id"],
                    "size": file["size"],
                },
            )
            documents.append(doc)

        print(f"✅ {len(documents)} Dokumente geladen")
        return documents


def get_sharepoint_documents_split(
    document_library_id: str = None,
    folder_path: str | None = None,
    recursive: bool = True
) -> List[Document]:
    """
    Lädt SharePoint-Dokumente mit Client Credentials Flow (Single Tenant).
    """

    print(f"DEBUG: Meine Tenant ID ist: '{settings.SP_TENANT_ID}'")
    print(f"DEBUG: Nutze Microsoft Graph API mit Client Credentials Flow")

    loader = SharePointDocumentLoader(
        tenant_id=settings.SP_TENANT_ID,
        client_id=settings.SP_CLIENT_ID,
        client_secret=settings.SP_CLIENT_SECRET,
        site_id=settings.SP_SITE_ID,
        drive_id=document_library_id or settings.SP_DRIVE_ID,
    )

    documents = loader.load_documents()
    chunked_docs = text_splitter.split_documents(documents)

    print(f"✅ {len(chunked_docs)} Chunks erstellt")

    return chunked_docs
